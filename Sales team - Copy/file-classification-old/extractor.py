"""
extractor.py - Multi-format text extraction.

Supports:
  - PDF  (pdfplumber with PyMuPDF fallback; OCR for scanned pages)
  - DOCX (python-docx)
  - XLSX (openpyxl)
  - TXT  (plain read)
  - Images: PNG, JPG, JPEG, TIFF (rostaing-ocr via SchemaOCRExtractor)
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Extracts plain text from documents of various formats.

    Args:
        ocr_enabled: When True, attempt OCR on images and scanned PDF pages.
    """

    SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
        {".pdf", ".docx", ".xlsx", ".txt", ".png", ".jpg", ".jpeg", ".tiff"}
    )

    def __init__(self, ocr_enabled: bool = True) -> None:
        self.ocr_enabled = ocr_enabled

    # ── Public API ─────────────────────────────────────────────────────────

    def extract(self, file_path: Path, max_pages: int = 3) -> tuple[str, str]:
        """
        Extract text from *file_path*.

        Args:
            file_path: Path to the document.
            max_pages: Maximum number of PDF pages to read. Limits token
                       usage when sending text to an LLM for scoring.
                       Set to 0 or None for no limit. Default is 3.

        Returns:
            A tuple of (extracted_text, error_message).
            *error_message* is an empty string on full success, or a
            human-readable description of any non-fatal problem encountered.
        """
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._extract_pdf(file_path, max_pages=max_pages or 0)
            elif suffix == ".docx":
                return self._extract_docx(file_path), ""
            elif suffix == ".xlsx":
                return self._extract_xlsx(file_path), ""
            elif suffix == ".txt":
                return self._extract_txt(file_path), ""
            elif suffix in {".png", ".jpg", ".jpeg", ".tiff"}:
                return self._extract_image(file_path), ""
            else:
                return "", f"Unsupported file type: {suffix}"
        except Exception as exc:
            logger.error("Extraction failed for %s: %s", file_path.name, exc, exc_info=True)
            return "", str(exc)

    # ── PDF ────────────────────────────────────────────────────────────────

    def _extract_pdf(self, file_path: Path, max_pages: int = 3) -> tuple[str, str]:
        """
        Try pdfplumber first; fall back to PyMuPDF; OCR if still empty.

        Args:
            max_pages: Read at most this many pages. 0 = unlimited.
        """
        text = ""
        error = ""

        page_limit_msg = f" (first {max_pages} pages)" if max_pages else ""
        logger.debug("Extracting PDF%s: %s", page_limit_msg, file_path.name)

        # Attempt 1 - pdfplumber
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    if max_pages and i >= max_pages:
                        break
                    page_text = page.extract_text() or ""
                    pages_text.append(page_text)
                text = "\n".join(pages_text)
        except Exception as exc:
            error = f"pdfplumber failed ({exc})"
            logger.debug("pdfplumber failed for %s: %s", file_path.name, exc)

        # Attempt 2 - PyMuPDF fallback
        if not text.strip():
            try:
                import fitz  # type: ignore  # PyMuPDF

                doc = fitz.open(str(file_path))
                pages = list(doc)[:max_pages] if max_pages else list(doc)
                pages_text = [page.get_text() for page in pages]
                doc.close()
                text = "\n".join(pages_text)
                error = ""
            except Exception as exc:
                error = f"{error}; fitz also failed ({exc})"
                logger.debug("PyMuPDF failed for %s: %s", file_path.name, exc)

        # Attempt 3 - OCR (scanned PDF -> rasterise each page)
        if not text.strip() and self.ocr_enabled:
            try:
                text, ocr_err = self._ocr_pdf(file_path, max_pages=max_pages)
                if ocr_err:
                    error = f"{error}; OCR: {ocr_err}".lstrip("; ")
                else:
                    error = ""
            except Exception as exc:
                error = f"{error}; OCR failed ({exc})".lstrip("; ")

        return text, error

    def _ocr_pdf(self, file_path: Path, max_pages: int = 3) -> tuple[str, str]:
        """Extract text from a scanned PDF using rostaing-ocr (via SchemaOCRExtractor).

        Args:
            max_pages: Read at most this many pages. 0 = unlimited.
        """
        try:
            from schema_ocr import SchemaOCRExtractor  # type: ignore
        except ImportError:
            return "", "schema_ocr / rostaing-ocr not installed; OCR unavailable."

        try:
            extractor = SchemaOCRExtractor(file_path)
            text = extractor.extract_layout_text(save_debug_output=False)
            return text, ""
        except Exception as exc:
            return "", f"rostaing-ocr extraction failed: {exc}"

    # ── DOCX ───────────────────────────────────────────────────────────────

    def _extract_docx(self, file_path: Path) -> str:
        """Extract paragraph text from a Word document."""
        from docx import Document  # type: ignore

        doc = Document(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)

    # ── XLSX ───────────────────────────────────────────────────────────────

    def _extract_xlsx(self, file_path: Path) -> str:
        """Read all cell values from all worksheets."""
        import openpyxl  # type: ignore

        wb = openpyxl.load_workbook(str(file_path), data_only=True, read_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    lines.append(row_text)
        wb.close()
        return "\n".join(lines)

    # ── TXT ────────────────────────────────────────────────────────────────

    def _extract_txt(self, file_path: Path) -> str:
        """Read a plain-text file, trying UTF-8 then latin-1 as a fallback."""
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return file_path.read_text(encoding="latin-1")

    # ── Images ─────────────────────────────────────────────────────────────

    def _extract_image(self, file_path: Path) -> str:
        """Run rostaing-ocr (via SchemaOCRExtractor) on a standalone image file."""
        if not self.ocr_enabled:
            return ""

        try:
            from schema_ocr import SchemaOCRExtractor  # type: ignore
        except ImportError:
            logger.warning("schema_ocr / rostaing-ocr not installed. Skipping OCR for %s.", file_path.name)
            return ""

        try:
            extractor = SchemaOCRExtractor(file_path)
            return extractor.extract_layout_text(save_debug_output=False)
        except Exception as exc:
            logger.warning("rostaing-ocr failed for %s: %s", file_path.name, exc)
            return ""
